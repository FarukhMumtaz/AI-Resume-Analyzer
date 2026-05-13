"""
resume_parser.py — Robust resume text extractor
Supports: PDF, DOCX, TXT
"""

import io
import re


def clean_text(text: str) -> str:
    """Clean extracted text: remove excess whitespace and blank lines."""
    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    # Normalize newlines — collapse 3+ blank lines to 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Strip leading/trailing whitespace
    text = text.strip()
    return text


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text)
        return clean_text("\n".join(pages_text))
    except ImportError:
        raise RuntimeError("pypdf library not installed. Run: pip install pypdf")
    except Exception as e:
        raise RuntimeError(f"PDF parsing failed: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return clean_text("\n".join(paragraphs))
    except ImportError:
        raise RuntimeError("python-docx library not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"DOCX parsing failed: {str(e)}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file."""
    try:
        # Try UTF-8 first, fall back to latin-1
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1")
        return clean_text(text)
    except Exception as e:
        raise RuntimeError(f"TXT parsing failed: {str(e)}")


def parse_resume(uploaded_file) -> str:
    """
    Main entry point. Accepts a Streamlit UploadedFile object.
    Returns extracted and cleaned text string.
    Raises RuntimeError with a user-friendly message on failure.
    """
    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif filename.endswith(".txt"):
        text = extract_text_from_txt(file_bytes)
    else:
        raise RuntimeError(f"Unsupported file type. Please upload a PDF, DOCX, or TXT file.")

    if not text or len(text.strip()) < 50:
        raise RuntimeError(
            "Could not extract enough text from this file. "
            "Please upload a readable PDF/DOCX/TXT resume. "
            "Image-based or scanned PDFs are not supported."
        )

    return text
